import sys
import json
import os
import re
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import spacy

# Intentar cargar spaCy para análisis gramatical
try:
    nlp = spacy.load("es_core_news_sm")
except:
    nlp = None

def get_word_class(text):
    """POS de la palabra analizada aislada (fallback cuando no hay contexto)."""
    if not nlp or not text.strip():
        return "Unknown"
    clean_text = re.sub(r'[^\w\s]', '', text)
    if not clean_text:
        return "Punct"
    doc = nlp(clean_text)
    if len(doc) > 0:
        return doc[0].pos_
    return "Unknown"

def get_hyperlinks(paragraph):
    """Extrae URLs de los elementos w:hyperlink de un párrafo."""
    urls = []
    try:
        # Paragraph -> Body -> Document
        doc = paragraph._parent._parent
        if not hasattr(doc, 'part'):
            return []
        rels = doc.part.rels
        for h in paragraph._element.xpath('.//w:hyperlink'):
            rId = h.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            if rId in rels:
                rel = rels[rId]
                if rel.reltype == RT.HYPERLINK:
                    urls.append(rel.target_ref)
    except Exception:
        pass
    return urls


def get_pos_at_offset(doc, char_offset):
    """Devuelve el pos_ del token que contiene char_offset en el doc de spaCy, o None."""
    t = get_token_at_offset(doc, char_offset)
    return t.pos_ if t is not None else None


def get_token_at_offset(doc, char_offset):
    """Devuelve el token que contiene char_offset en el doc de spaCy, o None (para lema, etc.)."""
    if not doc:
        return None
    for token in doc:
        if token.idx <= char_offset < token.idx + len(token.text):
            return token
    return None


def _lemma_looks_verb(lemma):
    """True si el lema tiene terminación de infinitivo español (-ar, -er, -ir)
    o si coincide con formas comunes mal etiquetadas.
    """
    if not lemma or len(lemma) < 2:
        return False
    low = lemma.lower()
    # Infinitivos
    if low.endswith("ar") or low.endswith("er") or low.endswith("ir"):
        return True
    return False

def _word_looks_like_verb(word):
    """Heurística para verbos conjugados que suelen fallar en POS tagging (PROPN).
    Enfocado en 1ra persona pretérito y formas comunes.
    """
    if not word or len(word) < 2:
        return False
    low = word.lower()
    # Terminaciones comunes de verbos conjugados que a veces se confunden con nombres
    # -í (nací, viví, crecí), -é (estudié, trabajé, empecé)
    if low.endswith("í") or low.endswith("é"):
        # Evitar nombres propios muy cortos que terminen así si es posible, 
        # pero para esta actividad el riesgo es bajo vs beneficio.
        return True
    # Terminaciones de pretérito imperfecto
    if low.endswith("aba") or low.endswith("ía"):
        return True
    return False

def pos_matches_category(pos_detected, category, token=None):
    """
    Comparación con equivalencias pedagógicas (UD vs. categorías de la actividad).
    - Verbo: VERB, AUX, y si spaCy etiquetó PROPN pero parece verbo.
    - Sustantivo: NOUN y PROPN.
    - Artículo: DET y PRON. Preposición: ADP y SCONJ.
    """
    if pos_detected == category:
        return True
    if category == "VERB":
        if pos_detected in ("AUX", "VERB"):
            return True
        # Caso especial: Misclassification de spaCy (PROPN para verbos conjugados)
        if pos_detected == "PROPN":
            # 1. Por lema (si lo detectó bien)
            if token is not None and _lemma_looks_verb(token.lemma_):
                return True
            # 2. Por terminación del texto (nací, viví, estudié)
            word = token.text if token is not None else ""
            if _word_looks_like_verb(word):
                return True
    if category == "NOUN" and pos_detected in ("NOUN", "PROPN"):
        return True
    if category == "ADJ" and pos_detected in ("ADJ", "DET", "PRON"):
        # En gramática escolar, los posesivos (mi, tu, su) se ven como adjetivos,
        # pero spaCy/UD los etiqueta como DET o PRON.
        return True
    if category == "DET" and pos_detected in ("DET", "PRON"):
        return True
    if category == "ADP" and pos_detected in ("ADP", "SCONJ"):
        return True
    return False

def hex_to_category(hex_color):
    """
    Mapea colores hexadecimales de Word a categorías gramaticales.
    Los colores pueden variar según el tema o selección manual.
    """
    if not hex_color: return None
    hex_color = hex_color.upper()
    
    # Verdes (Verbos)
    if hex_color in ['92D050', '00FF00', '76933C', 'C4D79B', 'EBF1DE', '00B050']:
        return "VERB"
    # Azules/Cianes (Sustantivos)
    if hex_color in ['0070C0', '00B0F0', '4F81BD', 'B8CCE4', 'DBE5F1', 'B7DEE8', '31859B', '92CDDC']:
        return "NOUN"
    # Amarillos (Adjetivos)
    if hex_color in ['FFFF00', 'FFC000', 'FFD966', 'FFF2CC', 'FCD5B4']:
        return "ADJ"
    # Rosas/Magenta (Adverbios)
    if hex_color in ['FF00FF', 'FF0066', 'E26B0A', 'F2DCDB', 'FDEADA', 'E46C0A']:
        return "ADV"
    # Grises (Artículos)
    if hex_color in ['A6A6A6', 'BFBFBF', 'D9D9D9', 'F2F2F2', '7F7F7F']:
        return "DET"
    # Morados (Preposiciones)
    if hex_color in ['7030A0', 'B1A0C7', 'E4DFEC', 'CCC1DA', '60497A']:
        return "ADP"
    
    return None

def get_run_color_category(run):
    # 1. Intentar con el resaltado estándar (Highlight)
    # WD_COLOR_INDEX: 4=Verde, 2=Azul, 7=Amarillo, 5=Rosa, 16=Gris, 12=Morado
    highlights = {
        4: "VERB", 11: "VERB",
        2: "NOUN", 9: "NOUN", 3: "NOUN", 10: "NOUN",
        7: "ADJ", 14: "ADJ",
        5: "ADV",
        16: "DET", 15: "DET",
        12: "ADP"
    }
    if run.font.highlight_color in highlights:
        return highlights[run.font.highlight_color]
    
    # 2. Intentar con el sombreado (Shading/Fill) - Muy común cuando se usa el bote de pintura
    try:
        rPr = run._element.get_or_add_rPr()
        shd = rPr.xpath('w:shd')
        if shd:
            fill = shd[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
            if fill and fill != 'auto':
                return hex_to_category(fill)
    except:
        pass
        
    return None

def process_docx(file_path):
    results = {
        "filename": os.path.basename(file_path),
        "filename_valid": False,
        "metrics": {
            "story_length": {"words": 0, "pages_approx": 0, "status": ""},
            "comparison_length": {"words": 0, "lines_approx": 0, "status": ""},
            "highlight_stats": {"total": 0, "correct": 0, "accuracy": 0},
            "audio_found": False
        },
        "segments": {
            "historia": {"text": "", "word_count": 0},
            "enlace": {"text": "", "url": ""},
            "highlights": [],
            "audio_url": "",
            "parrafo_comparacion": {"text": "", "line_count": 0, "word_count": 0},
            "transcripcion_audio": {"text": "", "error": None}
        },
        "errors": [],
        "titulo_presente": False,
        "fuente_arial_12": False,
        "evaluation": None
    }

    # 1. Validar nombre de archivo flexible
    # Debe contener Apellido_Nombre y alguna referencia a la actividad o grupo
    if "_" in results["filename"]:
        parts = results["filename"].split("_")
        if len(parts) >= 2:
            results["filename_valid"] = True
    elif "M2" in results["filename"] or "M02" in results["filename"]:
         results["filename_valid"] = True

    try:
        doc = Document(file_path)
    except Exception as e:
        results["errors"].append(f"Error al abrir DOCX: {str(e)}")
        return results

    cat_names = {
        "VERB": "Verbo (Verde)",
        "NOUN": "Sustantivo (Azul)",
        "ADJ": "Adjetivo (Amarillo)",
        "ADV": "Adverbio (Rosa)",
        "DET": "Artículo (Gris)",
        "ADP": "Preposición (Morado)"
    }

    # 4. Proceso de extracción: primero recorremos para ubicar secciones por posición
    all_highlights = []
    all_paragraphs = []  # (índice, texto, tiene_url, es_titulo_historia)
    url_pattern = re.compile(r'https?://[^\s<>"]+|www\.[^\s<>"]+')

    for idx, p in enumerate(doc.paragraphs):
        p_text = p.text.strip()
        if not p_text:
            all_paragraphs.append((idx, "", False, False))
            continue
        has_url = bool(url_pattern.search(p_text))
        # Título del relato: "El relato de mi historia" o variantes
        lower = p_text.lower()
        es_titulo = ("relato" in lower and "mi historia" in lower) or lower.strip() == "el relato de mi historia"
        all_paragraphs.append((idx, p_text, has_url, es_titulo))

    # Índices de corte: título de la historia, párrafo con enlace de audio
    titulo_idx = None
    enlace_idx = None
    reflection_keyword_idx = None

    for idx, (p_idx, p_text, has_url, es_titulo) in enumerate(all_paragraphs):
        para_obj = doc.paragraphs[p_idx]
        
        # 1. Buscar título si no existe
        if es_titulo and titulo_idx is None:
            titulo_idx = idx
            
        # 2. Buscar URLs (texto literal o hipervínculos)
        urls = url_pattern.findall(p_text)
        hyperlink_urls = get_hyperlinks(para_obj)
        all_urls = urls + hyperlink_urls
        
        if all_urls and enlace_idx is None and results["segments"]["audio_url"] == "":
            results["segments"]["audio_url"] = all_urls[0]
            results["metrics"]["audio_found"] = True
            enlace_idx = idx
            
        # 3. Buscar palabra clave "Reflexión" para robustez
        if "reflexión" in p_text.lower() and len(p_text.split()) < 10 and reflection_keyword_idx is None:
            reflection_keyword_idx = idx

    # Definir el punto de corte para la reflexión: 
    # Prioridad: Enlace de audio -> Palabra clave "Reflexión"
    cut_idx = enlace_idx if enlace_idx is not None else reflection_keyword_idx

    # Si no hay título explícito, considerar que la historia empieza en el primer párrafo sustancial
    # (evitando líneas de instrucciones)
    if titulo_idx is None:
        skip_start = ["actividad integradora", "mi historia", "propósito", "qué necesito", "qué entregaré", "cómo lo realizaré", "cómo se evaluará"]
        for idx, (_, p_text, has_url, _) in enumerate(all_paragraphs):
            if not p_text or has_url:
                continue
            lower = p_text.lower()
            if any(lower.startswith(k) or k in lower[:50] for k in skip_start):
                continue
            if len(p_text.split()) >= 20:  # párrafo largo = probable inicio del relato
                titulo_idx = idx - 1 if idx > 0 else -1
                break
        if titulo_idx is None:
            titulo_idx = -1

    # Título presente (para rúbrica)
    results["titulo_presente"] = titulo_idx is not None and titulo_idx >= 0

    # Fuente Arial 12 pt: muestreo en párrafos del cuerpo (historia)
    try:
        from docx.shared import Pt
        arial_runs = 0
        total_runs = 0
        start = (titulo_idx + 1) if (titulo_idx is not None and titulo_idx >= 0) else 0
        end = enlace_idx if enlace_idx is not None else len(doc.paragraphs)
        
        # Obtener valores por defecto del documento (Normal style y docDefaults)
        default_font_name = None
        default_font_size = None
        try:
            from docx.oxml.ns import qn
            # 1. Intentar desde docDefaults (XML)
            styles_element = doc.styles._element
            rFonts = styles_element.xpath('w:docDefaults/w:rPrDefault/w:rPr/w:rFonts')
            if rFonts:
                default_font_name = rFonts[0].get(qn('w:ascii'))
            
            sz = styles_element.xpath('w:docDefaults/w:rPrDefault/w:rPr/w:sz')
            if sz:
                val = sz[0].get(qn('w:val'))
                if val:
                    default_font_size = Pt(int(val) / 2) # sz is in half-points

            # 2. Fallback al estilo Normal si no hay docDefaults o para reforzar
            if not default_font_name or not default_font_size:
                normal_style = doc.styles['Normal']
                if not default_font_name: default_font_name = normal_style.font.name
                if not default_font_size: default_font_size = normal_style.font.size
        except:
            pass

        for idx in range(start, end):
            if idx >= len(doc.paragraphs):
                break
            para = doc.paragraphs[idx]
            
            # Valores del estilo del párrafo
            para_font_name = para.style.font.name if para.style and para.style.font else None
            para_font_size = para.style.font.size if para.style and para.style.font else None
            
            for run in para.runs:
                if not run.text.strip():
                    continue
                total_runs += 1
                
                # Resolución de fuente: Run -> Párrafo -> Documento
                name = run.font.name or para_font_name or default_font_name or ""
                name = name.lower()
                
                size = run.font.size or para_font_size or default_font_size
                size_ok = False
                if size is not None:
                    try:
                        # Si es un objeto Pt o similar, usar .pt
                        val_pt = size.pt if hasattr(size, 'pt') else size
                        size_ok = abs(val_pt - 12) < 0.5
                    except Exception:
                        size_ok = False
                
                if "arial" in name and size_ok:
                    arial_runs += 1
        
        results["fuente_arial_12"] = (total_runs > 0 and (arial_runs / total_runs) >= 0.5)
    except Exception:
        results["fuente_arial_12"] = False

    # Segmentos por posición: historia = entre título y enlace; reflexión = después del enlace
    story_content = []
    enlace_content = []
    reflection_content = []
    for idx, (_, p_text, has_url, es_titulo) in enumerate(all_paragraphs):
        if not p_text:
            continue
        if es_titulo:
            continue
        if cut_idx is not None and idx == cut_idx:
            enlace_content.append(p_text)
            # Si el párrafo del corte tiene mucha reflexión (más de una frase larga),
            # lo agregamos a reflexión también o solo lo tratamos como divisor
            continue
        if cut_idx is not None and idx > cut_idx:
            reflection_content.append(p_text)
            continue
        if titulo_idx is not None and titulo_idx >= 0 and idx > titulo_idx:
            story_content.append(p_text)
            continue
        # Sin título explícito: añadir párrafos que no parezcan instrucciones ni enlace
        if (titulo_idx is None or titulo_idx == -1) and (cut_idx is None or idx < cut_idx):
            if not has_url and not any(k in p_text.lower()[:60] for k in ["instrucciones", "rubrica", "propósito", "qué necesito", "qué entregaré", "cómo lo realizaré"]):
                story_content.append(p_text)

    # Segundo recorrido: extracción de colores (por párrafo), evaluando cada palabra EN CONTEXTO
    for idx, p in enumerate(doc.paragraphs):
        p_text = p.text  # sin strip para que los offsets coincidan con runs
        if not p_text.strip():
            continue
        # Análisis del párrafo completo para POS en contexto (mejor que palabra aislada)
        doc_para = nlp(p_text) if nlp else None
        run_start = 0
        for run in p.runs:
            category = get_run_color_category(run)
            if category:
                run_text = run.text
                if run_text:
                    words = re.findall(r'\b[\wáéíóúÁÉÍÓÚñÑ]+\b', run_text)
                    run_offset = 0
                    for w in words:
                        idx_in_run = run_text.find(w, run_offset)
                        if idx_in_run == -1:
                            break
                        word_start_in_para = run_start + idx_in_run
                        # POS en contexto; token permite lema para verbos conjugados mal etiquetados (PROPN→nacer)
                        token_at = get_token_at_offset(doc_para, word_start_in_para) if doc_para else None
                        if token_at is not None:
                            pos_detected = token_at.pos_
                        else:
                            pos_detected = get_word_class(w)
                        match = pos_matches_category(pos_detected, category, token=token_at)
                        all_highlights.append({
                            "word": w,
                            "pos_detected": pos_detected,
                            "expected_pos": category,
                            "match": bool(match),
                            "category": cat_names.get(category, "Otro")
                        })
                        run_offset = idx_in_run + len(w)
            run_start += len(run.text)

    # Resumen de Historia
    story_text = "\n\n".join(story_content)
    story_words = len(story_text.split())
    results["segments"]["historia"]["text"] = story_text
    results["segments"]["historia"]["word_count"] = story_words
    # 1-2 cuartillas (1 cuartilla ≈ 350 palabras)
    pages = round(story_words / 350, 1)
    results["metrics"]["story_length"]["words"] = story_words
    results["metrics"]["story_length"]["pages_approx"] = pages
    results["metrics"]["story_length"]["status"] = "Correcta" if 350 <= story_words <= 1000 else "Fuera de rango"

    # Resumen Enlace (párrafo con la URL del audio)
    enlace_text = "\n\n".join(enlace_content)
    results["segments"]["enlace"]["text"] = enlace_text
    results["segments"]["enlace"]["url"] = results["segments"]["audio_url"]

    # Resumen de Reflexión (párrafo 5-10 líneas sobre diferencias oral/escrito)
    refl_text = "\n\n".join(reflection_content)
    refl_words = len(refl_text.split())
    lines = max(1, refl_words // 12)  # Estimación de líneas (~12 palabras por línea)
    results["segments"]["parrafo_comparacion"]["text"] = refl_text
    results["segments"]["parrafo_comparacion"]["word_count"] = refl_words
    results["segments"]["parrafo_comparacion"]["line_count"] = lines
    results["metrics"]["comparison_length"]["words"] = refl_words
    results["metrics"]["comparison_length"]["lines_approx"] = lines
    results["metrics"]["comparison_length"]["status"] = "Correcta" if 5 <= lines <= 12 else "Revisar extensión"

    # Resumen de Highlights
    results["segments"]["highlights"] = all_highlights
    if all_highlights:
        correct = sum(1 for x in all_highlights if x["match"])
        results["metrics"]["highlight_stats"]["total"] = len(all_highlights)
        results["metrics"]["highlight_stats"]["correct"] = correct
        results["metrics"]["highlight_stats"]["accuracy"] = round((correct / len(all_highlights)) * 100, 1)

    # Texto completo para indicadores de IA y ortografía (historia + reflexión)
    full_text_for_analysis = (story_text + "\n\n" + refl_text).strip()
    try:
        from ai_indicators import detect_ai_indicators
        results["metrics"]["ai_indicadores"] = detect_ai_indicators(full_text_for_analysis)
    except Exception:
        results["metrics"]["ai_indicadores"] = None
    try:
        from orthography_check import count_orthography_errors
        count, err_reason, matches_list = count_orthography_errors(full_text_for_analysis)
        results["metrics"]["orthography_errors"] = count
        results["metrics"]["orthography_error_reason"] = err_reason
        results["metrics"]["orthography_matches"] = matches_list if matches_list is not None else None
    except Exception as e:
        results["metrics"]["orthography_errors"] = None
        results["metrics"]["orthography_error_reason"] = str(e)
        results["metrics"]["orthography_matches"] = None

    # Descarga y transcripción del audio (Drive, Dropbox, OneDrive, etc.)
    audio_duration_seconds = None
    if results["segments"]["audio_url"]:
        try:
            import tempfile
            script_dir = os.path.dirname(os.path.abspath(__file__))
            if script_dir not in sys.path:
                sys.path.insert(0, script_dir)
            from audio_transcribe import download_and_transcribe
            trans = download_and_transcribe(
                results["segments"]["audio_url"],
                dest_dir=tempfile.gettempdir(),
                language="es",
                model_size="base"
            )
            results["segments"]["transcripcion_audio"]["text"] = trans.get("text") or ""
            results["segments"]["transcripcion_audio"]["error"] = trans.get("error")
            audio_duration_seconds = trans.get("duration_seconds")
        except Exception as e:
            results["segments"]["transcripcion_audio"]["error"] = str(e)

    # Evaluación según rúbrica (métricas M2 AI2)
    try:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        if script_dir not in sys.path:
            sys.path.insert(0, script_dir)
        from evaluation_rubric import build_evaluation
        results["evaluation"] = build_evaluation(results, audio_duration_seconds=audio_duration_seconds)
    except Exception as e:
        results["evaluation"] = {"error": str(e)}

    return results

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"error": "No file path provided"}))
        sys.exit(1)
    
    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        print(json.dumps({"error": f"File not found: {file_path}"}))
        sys.exit(1)
        
    try:
        output = process_docx(file_path)
        print(json.dumps(output, ensure_ascii=False))
    except Exception as e:
        print(json.dumps({"error": str(e)}))
